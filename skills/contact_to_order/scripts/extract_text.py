#!/usr/bin/env python3
"""contact_to_order: local text extraction + Ollama LLM parse + JSON normalize.

Default pipeline (matches Dify DSL + model_routes.yaml):
  1. Classify attachment as image or document.
  2. For documents, extract plain text locally (PDF/DOCX/XLSX/...).
  3. Call Ollama OpenAI-compatible /v1/chat/completions with routed model/options.
  4. Clean model output via clean_contact_to_order_json.py rules.

Usage:
  python extract_text.py <attachment_path> [--issmple 0] [--callback-url URL]
  python extract_text.py <attachment_path> --extract-only

Hermes Agent: pass a on-disk path only (from chat-attachments index storage_path).
Never paste PDF/binary from chat into execute_code or Python strings.
"""
from __future__ import annotations

import argparse
import base64
import csv
import hashlib
import importlib.util
import json
import mimetypes
import os
import re
import sys
from datetime import datetime, timezone
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from urllib.parse import quote as url_quote

TEXT_EXTS = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".gif"}
DOCUMENT_EXTS = TEXT_EXTS | {".pdf", ".docx", ".xlsx", ".csv"}

SKILL_ROOT = Path(__file__).resolve().parent.parent
ROUTES_PATH = SKILL_ROOT / "model_routes.yaml"
PROMPT_PATH = SKILL_ROOT / "prompts" / "contact_to_order_prompt.md"
CLEAN_SCRIPT = Path(__file__).resolve().parent / "clean_contact_to_order_json.py"
DEFAULT_ISSMPLE = 0
HTTP_TIMEOUT_SEC = 600
MAX_WEB_URL_LEN = 20000
PDF_VISION_MAX_PAGES = 15
PDF_VISION_DPI = 144
WEB_URL_RUN_DIR = SKILL_ROOT / ".run"
WEB_URL_FILE = WEB_URL_RUN_DIR / "last_web_url.txt"
LLM_RAW_FILE = WEB_URL_RUN_DIR / "last_llm_raw.txt"


def emit(payload: dict[str, Any]) -> None:
    print(json.dumps(payload, ensure_ascii=False, indent=2))


def persist_web_url_file(web_url: str) -> tuple[str, str]:
    """Atomically write web_url; return (absolute path, ISO mtime UTC)."""
    WEB_URL_RUN_DIR.mkdir(parents=True, exist_ok=True)
    tmp = WEB_URL_FILE.with_suffix(".tmp")
    tmp.write_text(web_url, encoding="utf-8")
    tmp.replace(WEB_URL_FILE)
    mtime = datetime.fromtimestamp(WEB_URL_FILE.stat().st_mtime, tz=timezone.utc).isoformat()
    return str(WEB_URL_FILE.resolve()), mtime


def persist_llm_raw_file(raw: str) -> tuple[str, str]:
    """Atomically write Ollama raw content (pre-normalize); return (path, ISO mtime UTC)."""
    WEB_URL_RUN_DIR.mkdir(parents=True, exist_ok=True)
    tmp = LLM_RAW_FILE.with_suffix(".tmp")
    tmp.write_text(raw, encoding="utf-8")
    tmp.replace(LLM_RAW_FILE)
    mtime = datetime.fromtimestamp(LLM_RAW_FILE.stat().st_mtime, tz=timezone.utc).isoformat()
    return str(LLM_RAW_FILE.resolve()), mtime


def emit_for_agent(payload: dict[str, Any]) -> None:
    """Stdout for Hermes terminal — omit inline web_url (tool output often truncates long URLs)."""
    if payload.get("callback_applied") and payload.get("web_url_file"):
        slim: dict[str, Any] = {
            "ok": payload.get("ok", True),
            "file": payload.get("file"),
            "file_path": payload.get("file_path"),
            "parse_mode": payload.get("parse_mode"),
            "pdf_pages_sent": payload.get("pdf_pages_sent"),
            "file_type": payload.get("file_type"),
            "route": payload.get("route"),
            "callback_applied": True,
            "callback_host": payload.get("callback_host"),
            "web_url_file": payload.get("web_url_file"),
            "web_url_file_updated": payload.get("web_url_file_updated"),
            "web_url_file_mtime": payload.get("web_url_file_mtime"),
            "web_url_chars": payload.get("web_url_chars"),
            "llm_raw_output_file": payload.get("llm_raw_output_file"),
            "llm_raw_output_file_mtime": payload.get("llm_raw_output_file_mtime"),
            "llm_raw_output_chars": payload.get("llm_raw_output_chars"),
            "web_url_note": (
                "回复用户顺序：① Get-Content -Raw llm_raw_output_file 原样展示 LLM 原生输出；"
                "② 展示 file_path；③ Get-Content -Raw web_url_file 生成 Markdown 链接。"
                "禁止在聊天中粘贴 URL 正文（terminal 会截断）。"
            ),
        }
        print(json.dumps(slim, ensure_ascii=False, indent=2))
        return
    emit(payload)


def _default_routes() -> dict[str, Any]:
    return {
        "server": {
            "url": "http://192.168.70.249:11434/v1",
            "provider": "ollama",
            "api_format": "openai-compatible",
        },
        "models": {
            "image": {
                "0": {
                    "model": "qwen2.5vl:7b",
                    "num_ctx": 5048,
                    "num_predict": 5048,
                },
                "default": {
                    "model": "qwen3-vl:30b",
                    "num_ctx": 5048,
                    "num_predict": 5048,
                },
            },
            "document": {
                "0": {
                    "model": "qwen2.5vl:7b",
                    "num_ctx": 20480,
                    "num_predict": 10240,
                },
                "default": {
                    "model": "qwen3-vl:30b",
                    "num_ctx": 20480,
                    "num_predict": 10240,
                },
            },
        },
    }


def load_model_routes() -> dict[str, Any]:
    if not ROUTES_PATH.exists():
        return _default_routes()
    raw = ROUTES_PATH.read_text(encoding="utf-8")
    try:
        import yaml  # type: ignore

        data = yaml.safe_load(raw)
        return data if isinstance(data, dict) else _default_routes()
    except Exception:  # noqa: BLE001
        return _default_routes()


def select_route(routes: dict[str, Any], file_kind: str, issmple: int) -> dict[str, Any]:
    server = routes.get("server") or {}
    branch = (routes.get("models") or {}).get(file_kind) or {}
    route_key = "0" if issmple == 0 else "default"
    model_cfg = branch.get(route_key) or branch.get("default") or {}
    return {
        "base_url": str(server.get("url", "http://192.168.70.249:11434/v1")).rstrip("/"),
        "model": str(model_cfg.get("model", "qwen2.5vl:7b")),
        "num_ctx": int(model_cfg.get("num_ctx", 5048)),
        "num_predict": int(model_cfg.get("num_predict", 5048)),
        "file_kind": file_kind,
        "issmple": issmple,
        "route_key": route_key,
    }


def load_prompt_sections() -> tuple[str, str]:
    if not PROMPT_PATH.exists():
        raise FileNotFoundError(f"prompt not found: {PROMPT_PATH}")
    text = PROMPT_PATH.read_text(encoding="utf-8")
    system = ""
    user_rules = ""
    if "## System" in text:
        after_system = text.split("## System", 1)[1]
        if "## User Rules" in after_system:
            system, user_rules = after_system.split("## User Rules", 1)
        else:
            system = after_system
    system = system.strip()
    user_rules = user_rules.strip()
    return system, user_rules


def read_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_csv(path: Path) -> str:
    rows = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        for row in csv.reader(f):
            rows.append("\t".join(row))
    return "\n".join(rows)


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:  # noqa: BLE001
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError("PDF extraction requires pypdf or PyPDF2") from exc
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("DOCX extraction requires python-docx") from exc
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def read_xlsx(path: Path) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("XLSX extraction requires openpyxl") from exc
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            parts.append("\t".join("" if v is None else str(v) for v in row))
    return "\n".join(parts)


def classify_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in IMAGE_EXTS:
        return "image"
    if ext in DOCUMENT_EXTS or ext == ".pdf":
        return "document"
    return "unsupported"


def extract_document_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".csv":
        return read_csv(path)
    if ext in TEXT_EXTS:
        return read_plain(path)
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".xlsx":
        return read_xlsx(path)
    raise RuntimeError(f"unsupported document suffix: {ext}")


def image_data_url(path: Path) -> str:
    mime = mimetypes.guess_type(str(path))[0] or "image/jpeg"
    encoded = base64.standard_b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def render_pdf_pages(
    path: Path,
    *,
    max_pages: int = PDF_VISION_MAX_PAGES,
    dpi: int = PDF_VISION_DPI,
) -> list[Path]:
    """Render PDF pages to PNG for vision models (Ollama image_url does not accept PDF bytes)."""
    try:
        import fitz  # type: ignore  # PyMuPDF
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "PDF 视觉解析需要 PyMuPDF（pip install pymupdf）。"
            "扫描件/版式 PDF 请走此路径，勿仅依赖 pypdf 抽文本。"
        ) from exc

    key = hashlib.sha256(str(path.resolve()).encode("utf-8")).hexdigest()[:12]
    page_dir = WEB_URL_RUN_DIR / "pdf_pages" / key
    page_dir.mkdir(parents=True, exist_ok=True)

    doc = fitz.open(str(path))
    try:
        if doc.page_count == 0:
            raise RuntimeError("PDF 无页面")
        paths: list[Path] = []
        limit = min(doc.page_count, max_pages)
        for i in range(limit):
            pix = doc.load_page(i).get_pixmap(dpi=dpi)
            out = page_dir / f"page_{i + 1:03d}.png"
            pix.save(str(out))
            paths.append(out)
        if doc.page_count > max_pages:
            # marker file for debugging; vision prompt mentions truncated pages
            (page_dir / "_truncated.txt").write_text(
                f"total_pages={doc.page_count},sent={max_pages}", encoding="utf-8"
            )
        return paths
    finally:
        doc.close()


def call_ollama_chat(
    route: dict[str, Any],
    *,
    system: str,
    user_text: str,
    image_path: Path | None = None,
    image_paths: list[Path] | None = None,
) -> str:
    url = f"{route['base_url']}/chat/completions"
    vision_paths: list[Path] = []
    if image_paths:
        vision_paths = list(image_paths)
    elif image_path is not None:
        vision_paths = [image_path]

    if vision_paths:
        user_content: Any = [{"type": "text", "text": user_text}]
        for img in vision_paths:
            user_content.append(
                {"type": "image_url", "image_url": {"url": image_data_url(img)}}
            )
    else:
        user_content = user_text

    payload = {
        "model": route["model"],
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user_content},
        ],
        "stream": False,
        "options": {
            "num_ctx": route["num_ctx"],
            "num_predict": route["num_predict"],
        },
    }
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=HTTP_TIMEOUT_SEC) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"ollama HTTP {exc.code}: {detail[:2000]}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"ollama unreachable at {url}: {exc}") from exc

    choices = data.get("choices") or []
    if not choices:
        raise RuntimeError(f"ollama empty choices: {json.dumps(data, ensure_ascii=False)[:2000]}")
    message = choices[0].get("message") or {}
    content = message.get("content", "")
    if isinstance(content, list):
        parts = []
        for block in content:
            if isinstance(block, dict) and block.get("type") == "text":
                parts.append(str(block.get("text", "")))
        content = "\n".join(parts)
    if not str(content).strip():
        raise RuntimeError("ollama returned empty content")
    return str(content)


def load_clean_module():
    spec = importlib.util.spec_from_file_location("clean_contact_to_order_json", CLEAN_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load cleaner: {CLEAN_SCRIPT}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def normalize_model_output(raw: str) -> dict[str, Any]:
    clean = load_clean_module()
    obj = clean.parse_jsonish(raw)
    return clean.normalize_order(obj)


def is_valid_orderinfo(orderinfo: Any) -> bool:
    if not isinstance(orderinfo, list) or len(orderinfo) == 0:
        return False
    return all(isinstance(row, dict) for row in orderinfo)


def validate_callback_url(callback_url: str) -> str:
    url = callback_url.strip()
    url = url.replace("&amp;", "&").replace("&#38;", "&")
    if not url:
        raise ValueError("callbackURL 不能为空")
    if re.search(r"\s", url):
        raise ValueError(
            "callbackURL 含有空格（常见于 Windows PowerShell 未给 URL 加引号，"
            "& 被当作命令分隔符导致 URL 断裂，例如 :8080 8080/）。"
            "请用双引号包裹完整 URL，或使用 --callback-url-file / 环境变量 "
            "CONTACT_TO_ORDER_CALLBACK_URL。"
        )
    if re.search(r"[<>\"']", url):
        raise ValueError("callbackURL 含有非法字符 < > \" '")
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("callbackURL 必须使用 http 或 https")
    if not parsed.netloc:
        raise ValueError("callbackURL 缺少有效主机名")

    return url


def load_skill_params_file(file_path: str | None) -> dict[str, Any]:
    if not file_path:
        return {}
    raw = Path(file_path).expanduser().read_text(encoding="utf-8").strip()
    if not raw:
        return {}
    data = json.loads(raw)
    return data if isinstance(data, dict) else {}


def parse_issmple(value: Any, default: int = DEFAULT_ISSMPLE) -> int:
    if value is None:
        return default
    try:
        return int(str(value).strip())
    except ValueError:
        return default


def resolve_runtime_params(
    *,
    cli_issmple: int,
    skill_params: dict[str, Any],
    cli_callback: str | None,
    callback_file: str | None,
) -> tuple[int, str | None]:
    """Priority: [SkillParamsJSON] file > CLI callback flags > env; issmple from JSON if present."""
    issmple = (
        parse_issmple(skill_params.get("issmple"), cli_issmple)
        if skill_params
        else cli_issmple
    )
    callback: str | None = None
    if skill_params:
        raw = skill_params.get("callbackURL") or skill_params.get("callbackUrl")
        if isinstance(raw, str) and raw.strip():
            callback = raw.strip()
    if not callback:
        callback = load_callback_url(cli_callback, callback_file)
    return issmple, callback


def load_callback_url(
    cli_value: str | None,
    file_path: str | None,
) -> str | None:
    if file_path:
        text = Path(file_path).expanduser().read_text(encoding="utf-8").strip()
        if text:
            return text
    if cli_value and cli_value.strip():
        return cli_value.strip()
    env = os.environ.get("CONTACT_TO_ORDER_CALLBACK_URL", "").strip()
    return env or None


def encode_uri_component(value: str) -> str:
    """JavaScript encodeURIComponent 等价实现。"""
    return url_quote(value, safe="")


def sanitize_orderinfo_for_url(orderinfo: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Re-normalize so only schema fields are embedded in tempType (fixes partdesc%text etc.)."""
    clean = load_clean_module()
    return clean.normalize_order({"orderinfo": orderinfo})["orderinfo"]


def normalize_callback_base_for_append(base: str) -> str:
    """Dify 规则：callbackURL 以空 tempType= 结尾时 WebUrl = callbackURL + jsonData。"""
    trimmed = base.rstrip()
    if re.search(r"tempType=[^&\s].", trimmed, re.IGNORECASE):
        raise ValueError(
            "callbackURL 的 tempType= 后已有非空内容；请改为以空 tempType= 结尾 "
            "(例如 ...addSoDesktop&tempType=)，由脚本拼接 encodeURIComponent(JSON)。"
        )
    if trimmed.endswith("tempType=") or re.search(r"tempType=$", trimmed, re.IGNORECASE):
        return trimmed
    sep = "&" if "?" in trimmed else "?"
    return f"{trimmed}{sep}tempType="


def build_web_url(callback_url: str, orderinfo: list[dict[str, Any]]) -> str:
    base = validate_callback_url(callback_url)
    join_base = normalize_callback_base_for_append(base)
    rows = sanitize_orderinfo_for_url(orderinfo)
    post_data = {"tempType": rows}
    json_str = json.dumps(post_data, ensure_ascii=False, separators=(",", ":"))
    json_data = encode_uri_component(json_str)
    web_url = f"{join_base}{json_data}"
    if "&amp;" in web_url:
        raise ValueError("WebUrl 内部错误：不得包含 HTML 实体 &amp;")
    parsed = urlparse(web_url)
    base_parsed = urlparse(base)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise ValueError("拼接后的 WebUrl 无效")
    if parsed.netloc != base_parsed.netloc:
        raise ValueError(
            f"WebUrl 主机与 callbackURL 不一致: {parsed.netloc} != {base_parsed.netloc}"
        )
    if not web_url.startswith(join_base):
        raise ValueError("WebUrl 与 tempType= 拼接前缀不一致")
    if not web_url.startswith(base.rstrip()):
        raise ValueError("WebUrl 前缀与 callbackURL 不一致，禁止手工改写主机或 method 参数")
    if len(web_url) > MAX_WEB_URL_LEN:
        raise ValueError(f"WebUrl 过长（>{MAX_WEB_URL_LEN} 字符），请缩短 callbackURL 或减小报文")
    return web_url


def apply_callback(payload: dict[str, Any], callback_url: str | None) -> None:
    payload["web_url_file"] = str(WEB_URL_FILE.resolve())
    if not callback_url:
        payload["web_url_file_updated"] = False
        payload["web_url_file_skip_reason"] = "未提供 callbackURL，未写入 last_web_url.txt"
        return
    orderinfo = payload.get("orderinfo")
    if not is_valid_orderinfo(orderinfo):
        payload["callback_applied"] = False
        payload["callback_error"] = "orderinfo 无效或为空，未生成 WebUrl"
        payload["web_url_file_updated"] = False
        payload["web_url_file_skip_reason"] = "orderinfo 无效，保留上次 last_web_url.txt 内容"
        return
    try:
        rows = sanitize_orderinfo_for_url(orderinfo)
        web_url = build_web_url(callback_url, rows)
        url_file, url_mtime = persist_web_url_file(web_url)
        payload["callback_applied"] = True
        payload["callback_host"] = urlparse(web_url).netloc
        payload["web_url_file"] = url_file
        payload["web_url_file_updated"] = True
        payload["web_url_file_mtime"] = url_mtime
        payload["web_url_chars"] = len(web_url)
        payload["post_data"] = {"tempType": rows}
        payload["orderinfo"] = rows
        payload["web_url_note"] = (
            "完整 URL 已写入 web_url_file（见 web_url_file_mtime）。"
            "用 Get-Content -Raw web_url_file 读取后生成 [打开订单导入页面](url)；"
            "不要从 Markdown/HTML 复制（会出现 &amp;），不要手工拼 URL。"
        )
    except ValueError as exc:
        payload["callback_applied"] = False
        payload["callback_error"] = str(exc)
        payload["web_url_file_updated"] = False
        payload["web_url_file_skip_reason"] = (
            f"回调生成失败，未覆盖 last_web_url.txt: {exc}"
        )


def run_extract_only(path: Path) -> None:
    file_kind = classify_file(path)
    mime = mimetypes.guess_type(str(path))[0]
    ext = path.suffix.lower()
    if file_kind == "image":
        emit({
            "ok": True,
            "mode": "extract_only",
            "file_type": "image",
            "mime": mime,
            "text": "",
            "note": "Image text is not extracted locally; use default mode for Ollama vision parse.",
        })
        return
    if file_kind == "unsupported":
        emit({"ok": False, "error": "unsupported_file_type", "suffix": ext, "mime": mime})
        return
    text = extract_document_text(path)
    emit({
        "ok": True,
        "mode": "extract_only",
        "file_type": "document",
        "mime": mime,
        "chars": len(text),
        "text": text,
    })


def run_llm_parse(path: Path, issmple: int, callback_url: str | None = None) -> None:
    file_kind = classify_file(path)
    mime = mimetypes.guess_type(str(path))[0]
    ext = path.suffix.lower()
    if file_kind == "unsupported":
        emit({"ok": False, "error": "unsupported_file_type", "suffix": ext, "mime": mime})
        return

    routes = load_model_routes()
    route = select_route(routes, file_kind, issmple)
    system_prompt, user_rules = load_prompt_sections()

    parse_mode = "image" if file_kind == "image" else "document_text"
    document_text = ""
    pdf_page_images: list[Path] | None = None

    if ext == ".pdf":
        try:
            pdf_page_images = render_pdf_pages(path)
            parse_mode = "pdf_vision"
        except Exception as vision_exc:  # noqa: BLE001
            try:
                document_text = extract_document_text(path)
            except Exception as text_exc:  # noqa: BLE001
                emit({
                    "ok": False,
                    "error": "pdf_parse_failed",
                    "message": (
                        f"PDF 视觉渲染失败: {vision_exc}; "
                        f"文本抽取失败: {text_exc}"
                    ),
                    "file": str(path),
                    "hint": "pip install pymupdf，或提供可搜索文本层的 PDF",
                })
                return
            if not document_text.strip():
                emit({
                    "ok": False,
                    "error": "empty_document_text",
                    "message": (
                        f"PDF 视觉渲染失败: {vision_exc}; "
                        "文本层为空（多为扫描件）。请安装 pymupdf 后重试。"
                    ),
                    "file": str(path),
                })
                return
            parse_mode = "pdf_text_fallback"
    elif file_kind == "document":
        try:
            document_text = extract_document_text(path)
        except Exception as exc:  # noqa: BLE001
            emit({
                "ok": False,
                "error": "extract_failed",
                "message": str(exc),
                "file": str(path),
            })
            return
        if not document_text.strip():
            emit({
                "ok": False,
                "error": "empty_document_text",
                "message": "No text extracted; convert to searchable PDF/TXT or use OCR.",
                "file": str(path),
            })
            return

    if parse_mode == "pdf_vision" and pdf_page_images:
        user_body = (
            f"{user_rules}\n\n"
            f"附件为采购合同/订单 PDF，已按页转为图片（共 {len(pdf_page_images)} 页）。"
            "请仅根据这些页面图像内容提取订单 JSON，勿臆造字段。"
        )
        vision_images = pdf_page_images
        vision_single: Path | None = None
    elif file_kind == "document":
        user_body = (
            f"{user_rules}\n\n"
            "以下是从采购合同/订单附件中抽取的文本，请仅根据该文本提取订单 JSON：\n\n"
            f"{document_text}"
        )
        vision_images = None
        vision_single = None
    else:
        user_body = f"{user_rules}\n\n请根据附件图片提取订单 JSON。"
        vision_images = None
        vision_single = path

    try:
        raw = call_ollama_chat(
            route,
            system=system_prompt,
            user_text=user_body,
            image_path=vision_single,
            image_paths=vision_images,
        )
        llm_raw_file, llm_raw_mtime = persist_llm_raw_file(raw)
        normalized = normalize_model_output(raw)
    except Exception as exc:  # noqa: BLE001
        emit({
            "ok": False,
            "error": "llm_parse_failed",
            "message": str(exc),
            "file": str(path),
            "route": route,
        })
        return

    payload: dict[str, Any] = {
        "ok": True,
        "file": path.name,
        "file_path": str(path),
        "file_type": file_kind,
        "parse_mode": parse_mode,
        "pdf_pages_sent": len(pdf_page_images) if pdf_page_images else 0,
        "mime": mime,
        "route": route,
        "llm_raw_output_file": llm_raw_file,
        "llm_raw_output_file_mtime": llm_raw_mtime,
        "llm_raw_output_chars": len(raw),
        **normalized,
    }
    apply_callback(payload, callback_url)
    emit_for_agent(payload)


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="contact_to_order Ollama contract parser")
    parser.add_argument("file_path", help="Path to contract attachment")
    parser.add_argument(
        "--issmple",
        type=int,
        default=DEFAULT_ISSMPLE,
        help="0 -> qwen2.5vl route; non-zero -> qwen3-vl route (Dify compat name)",
    )
    parser.add_argument(
        "--extract-only",
        action="store_true",
        help="Only run local text extraction (no Ollama call)",
    )
    parser.add_argument(
        "--callback-url",
        "--callbackURL",
        dest="callback_url",
        default=None,
        metavar="URL",
        help="Callback base URL (must be quoted in PowerShell if URL contains &)",
    )
    parser.add_argument(
        "--callback-url-file",
        dest="callback_url_file",
        default=None,
        metavar="PATH",
        help="Read callback URL from a one-line text file (avoids shell & escaping)",
    )
    parser.add_argument(
        "--skill-params-file",
        dest="skill_params_file",
        default=None,
        metavar="PATH",
        help="JSON file from Hermes [SkillParamsJSON] block (callbackURL, issmple, ...)",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> None:
    args = parse_args(argv or sys.argv[1:])
    path = Path(args.file_path).expanduser().resolve()
    if not path.exists():
        emit({"ok": False, "error": "file_not_found", "file": str(path)})
        return
    try:
        skill_params = load_skill_params_file(args.skill_params_file)
    except json.JSONDecodeError as exc:
        emit({"ok": False, "error": "invalid_skill_params_json", "message": str(exc)})
        return
    issmple, callback_url = resolve_runtime_params(
        cli_issmple=args.issmple,
        skill_params=skill_params,
        cli_callback=args.callback_url,
        callback_file=args.callback_url_file,
    )
    if args.extract_only:
        run_extract_only(path)
        return
    run_llm_parse(path, issmple, callback_url=callback_url)


if __name__ == "__main__":
    main()
